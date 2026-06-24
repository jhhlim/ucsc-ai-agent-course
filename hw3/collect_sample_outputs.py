#!/usr/bin/env python3
"""Run README sample prompts through each ADK agent and save outputs to DOCX."""

from __future__ import annotations

import asyncio
import json
import sys
from datetime import datetime, timezone
from pathlib import Path

_HW3 = Path(__file__).resolve().parent
if str(_HW3) not in sys.path:
    sys.path.insert(0, str(_HW3))

from google.adk.runners import Runner
from google.adk.sessions import InMemorySessionService
from google.genai.types import Content, Part

from mortgage_agents.builders import (
    affordability_analyzer_agent,
    compliance_critic_agent,
    payment_calculator_agent,
    rate_finder_agent,
    supervisor_agent,
)

SAMPLES: list[tuple[str, str, object]] = [
    (
        "mortgage_supervisor — full pipeline",
        "mortgage_supervisor_full",
        supervisor_agent,
        """I'm looking at a home in Berryessa, San Jose for $1,280,000.
20% down, 30-year fixed at 6%.
Monthly costs: property tax $1,300, home insurance $150, HOA $25, utilities $300.
Household income $190,000/year, credit score 770, no other monthly debts.
Calculate total monthly payment, LTV, PMI, cash to close, and DTI.
Explain tradeoffs — educational only, no loan recommendation.""",
    ),
    (
        "mortgage_supervisor — compare down payments",
        "mortgage_supervisor_compare",
        supervisor_agent,
        """Compare 20% down vs 10% down on a $1,280,000 Berryessa home at 6% (30-year fixed).
Same costs: tax $1,300, insurance $150, HOA $25, utilities $300.
Income $190,000, credit 770. Show monthly payment, PMI, cash to close, and DTI side by side.
Educational only — no recommendation.""",
    ),
    (
        "rate_finder",
        "rate_finder",
        rate_finder_agent,
        """For a conventional 30-year fixed loan on a $1,280,000 Berryessa (San Jose) purchase,
explain how today's national 30-year benchmark rate (FRED MORTGAGE30US) compares to
a 6% quoted rate. Also explain 30-year fixed vs 15-year fixed vs 5/1 ARM tradeoffs
for this price range. No recommendation — just facts.""",
    ),
    (
        "payment_calculator",
        "payment_calculator",
        payment_calculator_agent,
        """Purchase price: $1,280,000
Down payment: 20% ($256,000)
Interest rate: 6%, 30-year fixed
Property tax: $1,300/mo
Home insurance: $150/mo
HOA: $25/mo
Utilities: $300/mo

Calculate loan amount, LTV, monthly P&I, PMI (if any), total monthly housing cost,
and estimated cash to close (include ~3% closing costs).""",
    ),
    (
        "affordability_analyzer",
        "affordability_analyzer",
        affordability_analyzer_agent,
        """Gross annual household income: $190,000
Credit score: 770
Other monthly debts: $0
Proposed total monthly housing payment: $7,914
(includes P&I ~$6,139, tax $1,300, insurance $150, HOA $25, utilities $300)

What credit tier does 770 map to? What are front-end and back-end DTI ratios?
How do those compare to typical conventional guidelines? Illustrative only.""",
    ),
    (
        "compliance_critic",
        "compliance_critic",
        compliance_critic_agent,
        '''Review this draft mortgage summary for math errors and advice-giving language:

"On a $1,280,000 Berryessa home with $256,000 down (20%), the loan is $1,024,000
at 6% for 30 years. Monthly P&I is $6,139. With tax ($1,300), insurance ($150),
HOA ($25), and utilities ($300), total monthly cost is $7,914.
At $190,000 income, front-end DTI is 50%. You should take this loan — you can afford it."

Verify P&I with loan_amount=1024000, rate=6%, term=30, reported_monthly_pi=6139.
Flag any issues and rewrite without advice language.''',
    ),
]


async def run_one(title: str, app_name: str, agent: object, prompt: str) -> dict:
    print(f"\n=== Running: {title} ===", flush=True)
    session_service = InMemorySessionService()
    session = await session_service.create_session(
        app_name=app_name,
        user_id="sample_outputs",
        session_id=f"{app_name}_{datetime.now(timezone.utc).strftime('%H%M%S')}",
    )
    runner = Runner(agent=agent, app_name=app_name, session_service=session_service)
    message = Content(role="user", parts=[Part(text=prompt)])
    events = runner.run_async(
        user_id="sample_outputs",
        session_id=session.id,
        new_message=message,
    )

    final_text = ""
    try:
        async for event in events:
            if not event.is_final_response():
                continue
            if not event.content or not event.content.parts:
                continue
            chunks = [p.text for p in event.content.parts if getattr(p, "text", None)]
            if chunks:
                final_text = "\n".join(chunks)
    finally:
        await events.aclose()

    print(f"Done: {title} ({len(final_text)} chars)", flush=True)
    return {
        "title": title,
        "agent": app_name,
        "prompt": prompt.strip(),
        "response": final_text.strip() or "(No final response captured)",
    }


async def main() -> list[dict]:
    results = []
    for title, app_name, agent, prompt in SAMPLES:
        try:
            results.append(await run_one(title, app_name, agent, prompt))
        except Exception as exc:
            results.append(
                {
                    "title": title,
                    "agent": app_name,
                    "prompt": prompt.strip(),
                    "response": f"ERROR: {exc}",
                }
            )
    return results


def write_docx(results: list[dict], path: Path) -> None:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading("HW3 Mortgage Agent — Sample Prompt Outputs", level=0)
    doc.add_paragraph(
        f"Berryessa $1.28M scenario · Generated {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}"
    )
    doc.add_paragraph(
        "Educational outputs from Doubleword/LiteLLM via Google ADK. "
        "Not financial advice."
    )

    for i, item in enumerate(results, start=1):
        doc.add_heading(f"{i}. {item['title']}", level=1)
        p = doc.add_paragraph()
        p.add_run("Agent: ").bold = True
        p.add_run(item["agent"])

        doc.add_heading("Prompt", level=2)
        for line in item["prompt"].splitlines():
            doc.add_paragraph(line, style="Intense Quote")

        doc.add_heading("Output", level=2)
        for line in item["response"].splitlines():
            para = doc.add_paragraph(line)
            para.paragraph_format.space_after = Pt(2)

        doc.add_page_break()

    doc.save(path)


if __name__ == "__main__":
    out_json = _HW3 / "sample_prompt_outputs.json"
    out_docx = _HW3 / "HW3_Sample_Prompt_Outputs.docx"

    results = asyncio.run(main())
    out_json.write_text(json.dumps(results, indent=2, ensure_ascii=False), encoding="utf-8")
    write_docx(results, out_docx)
    print(f"\nWrote {out_docx} and {out_json}")
